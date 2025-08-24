import io
import os

import subprocess
import tempfile
import zipfile
import asyncio
import logging
import time
from pathlib import Path
from typing import List, Optional, Dict, Union
from datetime import datetime

from fastapi import FastAPI, File, Form, UploadFile, Request, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response, StreamingResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from pydantic import BaseModel, Field

from pypdf import PdfReader, PdfWriter
import pikepdf

from PIL import Image
import img2pdf
from pdf2image import convert_from_bytes
from docx import Document


APP_ROOT = Path(__file__).resolve().parent.parent

# Production environment detection
IS_PRODUCTION = os.getenv("ENVIRONMENT", "development").lower() == "production"

app = FastAPI(
    title="PDF Tools API",
    description="Professional PDF conversion and manipulation tools",
    version="2.0.0",
    docs_url="/docs" if not IS_PRODUCTION else None,  # Enable docs in development
    redoc_url="/redoc" if not IS_PRODUCTION else None   # Enable redoc in development
)

# Rate limiting configuration - Production optimized
limiter = Limiter(key_func=get_remote_address)
RATE_LIMIT_PER_MINUTE = int(os.getenv("APP_RATE_LIMIT_PER_MINUTE", "60" if IS_PRODUCTION else "30"))
RATE_LIMIT_PER_HOUR = int(os.getenv("APP_RATE_LIMIT_PER_HOUR", "500" if IS_PRODUCTION else "300"))

# CORS configuration - Production optimized
if IS_PRODUCTION:
    # In production, only allow specific domains
    ALLOWED_ORIGINS = os.getenv("APP_ALLOWED_ORIGINS", "").split(",")
    if not ALLOWED_ORIGINS or ALLOWED_ORIGINS == [""]:
        ALLOWED_ORIGINS = ["*"]  # Fallback for production
else:
    # Development allows localhost
    ALLOWED_ORIGINS = os.getenv("APP_ALLOWED_ORIGINS", "http://localhost:3000,http://localhost:2000").split(",")

# Service start time for uptime calculation
SERVICE_START_TIME = time.time()

# Rate limiting middleware
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# Enhanced CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["X-RateLimit-Limit", "X-RateLimit-Remaining", "X-RateLimit-Reset"],
    max_age=3600,
)

# Serve static files from the static directory
static_dir = str(APP_ROOT / "static")
app.mount("/static", StaticFiles(directory=static_dir), name="static")

# ----------------------------------------------------------------------------
# Security/limits configuration - Production optimized
# ----------------------------------------------------------------------------
MAX_UPLOAD_MB = int(os.getenv("APP_MAX_UPLOAD_MB", "50" if IS_PRODUCTION else "20"))  # Higher limit for production
MAX_UPLOAD_BYTES = MAX_UPLOAD_MB * 1024 * 1024
SUBPROC_TIMEOUT = int(os.getenv("APP_SUBPROCESS_TIMEOUT", "900" if IS_PRODUCTION else "600"))  # Higher timeout for production
MAX_CONCURRENCY = int(os.getenv("APP_MAX_CONCURRENCY", "8" if IS_PRODUCTION else "4"))  # Higher concurrency for production
SEMAPHORE = asyncio.Semaphore(MAX_CONCURRENCY)

# Configure logging - Production optimized
if IS_PRODUCTION:
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
            logging.StreamHandler(),
            logging.FileHandler('/var/log/pdf-tools-api.log')
        ]
    )
else:
    logging.basicConfig(level=logging.INFO)

logger = logging.getLogger(__name__)

# Response models
class ErrorResponse(BaseModel):
    error: str = Field(..., description="Error message")
    code: str = Field(..., description="Error code")
    timestamp: datetime = Field(default_factory=datetime.utcnow)
    details: Optional[Dict[str, str]] = Field(None, description="Additional error details")

class SuccessResponse(BaseModel):
    message: str = Field(..., description="Success message")
    timestamp: datetime = Field(default_factory=datetime.utcnow)
    data: Optional[Dict[str, str]] = Field(None, description="Response data")

class HealthResponse(BaseModel):
    status: str = Field(..., description="Service status")
    timestamp: datetime = Field(default_factory=datetime.utcnow)
    version: str = Field(..., description="API version")
    uptime: float = Field(..., description="Service uptime in seconds")
    memory_usage: Dict[str, Union[str, int, float]] = Field(..., description="Memory usage information")
    environment: str = Field(..., description="Current environment")


def _ensure_size_limit(size: int) -> None:
    if size > MAX_UPLOAD_BYTES:
        logger.warning(f"File size {size} bytes exceeds limit {MAX_UPLOAD_BYTES}")
        raise ValueError(f"File too large. Limit is {MAX_UPLOAD_MB} MB")


def _is_pdf(content: bytes) -> bool:
    return content.startswith(b"%PDF")


def _is_image(content: bytes) -> bool:
    try:
        img = Image.open(io.BytesIO(content))
        img.verify()  # type: ignore[attr-defined]
        return True
    except Exception:
        return False


def _is_docx(content: bytes, filename: str) -> bool:
    # Very light check: ZIP magic and .docx extension
    return filename.lower().endswith(".docx") and content.startswith(b"PK")


@app.middleware("http")
async def security_headers(request: Request, call_next):
    # Enforce basic size check via Content-Length if present
    try:
        cl_header = request.headers.get("content-length")
        if cl_header and cl_header.isdigit():
            size_mb = int(cl_header) / (1024 * 1024)
            logger.info(f"Request size: {size_mb:.1f} MB")
            if int(cl_header) > MAX_UPLOAD_BYTES:
                logger.warning(f"Request too large: {size_mb:.1f} MB > {MAX_UPLOAD_MB} MB")
                return JSONResponse(
                    {"error": f"Payload too large. Max {MAX_UPLOAD_MB} MB"}, status_code=413
                )
    except Exception:
        pass

    try:
        response = await call_next(request)
        # Security headers - Production enhanced
        response.headers.setdefault("X-Content-Type-Options", "nosniff")
        response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
        response.headers.setdefault("Permissions-Policy", "geolocation=(), camera=(), microphone=()")
        response.headers.setdefault("X-Frame-Options", "DENY")
        response.headers.setdefault("X-XSS-Protection", "1; mode=block")
        
        # Enhanced CSP for production - Local resources only
        if IS_PRODUCTION:
            response.headers.setdefault("Content-Security-Policy", 
                "default-src 'self'; "
                "script-src 'self' 'unsafe-inline'; "
                "style-src 'self' 'unsafe-inline'; "
                "img-src 'self' data:; "
                "font-src 'self' data:; "
                "connect-src 'self'; "
                "frame-ancestors 'none';"
            )
        else:
            # Development CSP - Local resources only
            response.headers.setdefault("Content-Security-Policy", 
                "default-src 'self'; "
                "script-src 'self' 'unsafe-inline'; "
                "style-src 'self' 'unsafe-inline'; "
                "img-src 'self' data:; "
                "font-src 'self' data:; "
                "connect-src 'self';"
            )
        
        return response
    except Exception as e:
        logger.error(f"Request processing error: {e}")
        return JSONResponse({"error": "Internal server error"}, status_code=500)


@app.get("/")
def root_index():
    index_path = APP_ROOT / "pdfedit.html"
    if index_path.exists():
        return FileResponse(str(index_path))
    # Fallback text if html missing
    return Response(content="PDF Tools API running.", media_type="text/plain")


@app.get("/about")
def about_page():
    page = APP_ROOT / "about.html"
    if page.exists():
        return FileResponse(str(page))
    return Response(content="About", media_type="text/plain")


@app.get("/contact")
def contact_page():
    page = APP_ROOT / "contact.html"
    if page.exists():
        return FileResponse(str(page))
    return Response(content="Contact", media_type="text/plain")


@app.get("/robots.txt")
def robots_txt(request: Request):
    scheme = request.url.scheme
    host = request.headers.get("host", "localhost:2000")
    base_url = f"{scheme}://{host}"
    content = (
        "User-agent: *\n"
        "Allow: /\n"
        "Disallow: /api/\n"
        f"Sitemap: {base_url}/sitemap.xml\n"
    )
    return Response(content=content, media_type="text/plain")


@app.get("/sitemap.xml")
def sitemap_xml(request: Request):
    scheme = request.url.scheme
    host = request.headers.get("host", "localhost:2000")
    base_url = f"{scheme}://{host}"
    urls = [
        {"loc": f"{base_url}/", "changefreq": "weekly", "priority": "1.0"},
        {"loc": f"{base_url}/about", "changefreq": "yearly", "priority": "0.6"},
        {"loc": f"{base_url}/contact", "changefreq": "yearly", "priority": "0.5"},
    ]
    items = "".join(
        [
            f"<url><loc>{u['loc']}</loc><changefreq>{u['changefreq']}</changefreq><priority>{u['priority']}</priority></url>"
            for u in urls
        ]
    )
    xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\"?>"
        "<urlset xmlns=\"http://www.sitemaps.org/schemas/sitemap/0.9\">"
        f"{items}"  # noqa: E231
        "</urlset>"
    )
    return Response(content=xml, media_type="application/xml")


# Health and monitoring endpoints
@app.get("/health", response_model=HealthResponse, tags=["Health"])
@limiter.limit(f"{RATE_LIMIT_PER_MINUTE}/minute")
async def health_check(request: Request):
    """Health check endpoint for monitoring"""
    try:
        import psutil
        memory = psutil.virtual_memory()
        uptime = time.time() - SERVICE_START_TIME
        
        return HealthResponse(
            status="healthy",
            version="2.0.0",
            uptime=uptime,
            memory_usage={
                "total": memory.total,
                "available": memory.available,
                "percent": memory.percent,
                "used": memory.used
            },
            environment=os.getenv("ENVIRONMENT", "development")
        )
    except ImportError:
        # Fallback if psutil not available
        uptime = time.time() - SERVICE_START_TIME
        return HealthResponse(
            status="healthy",
            version="2.0.0",
            uptime=uptime,
            memory_usage={"note": "psutil not available"},
            environment=os.getenv("ENVIRONMENT", "development")
        )


@app.get("/health/ready", tags=["Health"])
@limiter.limit(f"{RATE_LIMIT_PER_MINUTE}/minute")
async def readiness_check(request: Request):
    """Readiness check for Kubernetes/container orchestration"""
    return {"status": "ready", "timestamp": datetime.utcnow()}


@app.get("/health/live", tags=["Health"])
@limiter.limit(f"{RATE_LIMIT_PER_MINUTE}/minute")
async def liveness_check(request: Request):
    """Liveness check for Kubernetes/container orchestration"""
    return {"status": "alive", "timestamp": datetime.utcnow()}


@app.get("/metrics", tags=["Monitoring"])
@limiter.limit(f"{RATE_LIMIT_PER_MINUTE}/minute")
async def metrics_endpoint(request: Request):
    """Prometheus metrics endpoint"""
    try:
        import psutil
        memory = psutil.virtual_memory()
        disk = psutil.disk_usage('/')
        
        metrics = {
            "pdf_tools_memory_bytes_total": memory.total,
            "pdf_tools_memory_bytes_used": memory.used,
            "pdf_tools_memory_percent": memory.percent,
            "pdf_tools_disk_bytes_total": disk.total,
            "pdf_tools_disk_bytes_used": disk.used,
            "pdf_tools_uptime_seconds": time.time() - SERVICE_START_TIME
        }
        
        return Response(
            content="\n".join([f"{k} {v}" for k, v in metrics.items()]),
            media_type="text/plain"
        )
    except ImportError:
        return Response(
            content="pdf_tools_uptime_seconds " + str(time.time() - SERVICE_START_TIME),
            media_type="text/plain"
        )


def _stream_bytes(bytes_data: bytes, filename: str, media_type: str) -> StreamingResponse:
    return StreamingResponse(
        io.BytesIO(bytes_data),
        media_type=media_type,
        headers={
            "Content-Disposition": f"attachment; filename=\"{filename}\""
        },
    )


@app.post("/api/merge", tags=["PDF Operations"])
@limiter.limit(f"{RATE_LIMIT_PER_HOUR}/hour")
async def merge_pdfs(request: Request, files: List[UploadFile] = File(...)):
    async with SEMAPHORE:
        # Single file for single output, multiple files for zip
        if len(files) == 1:
            # Single file - direct merge
            content = await files[0].read()
            _ensure_size_limit(len(content))
            if not _is_pdf(content):
                return JSONResponse({"error": "Only PDF files are allowed"}, status_code=400)
            reader = PdfReader(io.BytesIO(content))
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)
            buf = io.BytesIO()
            writer.write(buf)
            writer.close()
            buf.seek(0)
            out_name = f"merged_{Path(files[0].filename or 'document.pdf').stem}.pdf"
            return _stream_bytes(buf.getvalue(), out_name, "application/pdf")
        else:
            # Multiple files - output as zip
            with tempfile.TemporaryDirectory() as td:
                out_zip = Path(td) / "merged_pdfs.zip"
                
                with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                    for i, f in enumerate(files):
                        try:
                            content = await f.read()
                            _ensure_size_limit(len(content))
                            if not _is_pdf(content):
                                continue  # Skip non-PDF files
                            
                            reader = PdfReader(io.BytesIO(content))
                            writer = PdfWriter()
                            for page in reader.pages:
                                writer.add_page(page)
                            
                            # Save as temporary file
                            temp_pdf = Path(td) / f"merged_{i+1:03d}.pdf"
                            with open(temp_pdf, "wb") as pf:
                                writer.write(pf)
                            writer.close()
                            
                            # Add to zip
                            original_name = Path(f.filename or f"document_{i+1}.pdf").stem
                            zf.write(temp_pdf, arcname=f"merged_{original_name}.pdf")
                            
                        except Exception as e:
                            logger.error(f"Error processing {f.filename}: {str(e)}")
                            continue
                
                # Return zip file
                zip_data = out_zip.read_bytes()
                return _stream_bytes(zip_data, "merged_pdfs.zip", "application/zip")


def _parse_ranges(ranges: str, num_pages: int) -> List[int]:
    selected: List[int] = []
    token_list = [t.strip() for t in ranges.split(",") if t.strip()]
    for token in token_list:
        if "-" in token:
            start_s, end_s = token.split("-", 1)
            start = max(1, int(start_s))
            end = min(num_pages, int(end_s))
            if start <= end:
                selected.extend(list(range(start, end + 1)))
        else:
            idx = int(token)
            if 1 <= idx <= num_pages:
                selected.append(idx)
    # unique and preserve order
    seen = set()
    ordered = []
    for p in selected:
        if p not in seen:
            ordered.append(p)
            seen.add(p)
    return ordered


@app.post("/api/split", tags=["PDF Operations"])
@limiter.limit(f"{RATE_LIMIT_PER_HOUR}/hour")
async def split_pdf(
    request: Request,
    file: UploadFile = File(...),
    pages: Optional[int] = Form(None),
    ranges: Optional[str] = Form(None),
):
    content = await file.read()
    _ensure_size_limit(len(content))
    if not _is_pdf(content):
        return JSONResponse({"error": "Only PDF files are allowed"}, status_code=400)
    reader = PdfReader(io.BytesIO(content))
    total_pages = len(reader.pages)
    
    if pages and pages > 0:
        # Split into chunks of specified size
        with tempfile.TemporaryDirectory() as td:
            out_zip = Path(td) / "split_pages.zip"
            with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                chunk_num = 1
                for i in range(0, total_pages, pages):
                    writer = PdfWriter()
                    end_idx = min(i + pages, total_pages)
                    for j in range(i, end_idx):
                        writer.add_page(reader.pages[j])
                    
                    chunk_path = Path(td) / f"chunk_{chunk_num:03d}.pdf"
                    with open(chunk_path, "wb") as pf:
                        writer.write(pf)
                    writer.close()
                    zf.write(chunk_path, arcname=chunk_path.name)
                    chunk_num += 1
            data = out_zip.read_bytes()
            return _stream_bytes(data, "split_pages.zip", "application/zip")
    
    elif ranges:
        # Split by specific page ranges
        page_numbers = _parse_ranges(ranges, total_pages)
        writer = PdfWriter()
        for p in page_numbers:
            writer.add_page(reader.pages[p - 1])
        buf = io.BytesIO()
        writer.write(buf)
        writer.close()
        buf.seek(0)
        return _stream_bytes(buf.getvalue(), "extracted.pdf", "application/pdf")
    
    else:
        # Default: split every page to a separate pdf and zip
        with tempfile.TemporaryDirectory() as td:
            out_zip = Path(td) / "split_pages.zip"
            with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                for i, page in enumerate(reader.pages, start=1):
                    writer = PdfWriter()
                    writer.add_page(page)
                    page_path = Path(td) / f"page_{i:03d}.pdf"
                    with open(page_path, "wb") as pf:
                        writer.write(pf)
                    writer.close()
                    zf.write(page_path, arcname=page_path.name)
            data = out_zip.read_bytes()
            return _stream_bytes(data, "split_pages.zip", "application/zip")


@app.post("/api/compress")
async def compress_pdf(
    file: UploadFile = File(...),
    preset: str = Form("medium"),
):
    preset_map = {
        "low": "/screen",
        "medium": "/ebook",
        "high": "/prepress",
    }
    gs_preset = preset_map.get(preset, "/ebook")
    with tempfile.TemporaryDirectory() as td:
        in_path = Path(td) / "input.pdf"
        out_path = Path(td) / "output.pdf"
        in_path.write_bytes(await file.read())
        cmd = [
            "gs",
            "-sDEVICE=pdfwrite",
            "-dCompatibilityLevel=1.4",
            f"-dPDFSETTINGS={gs_preset}",
            "-dNOPAUSE",
            "-dQUIET",
            "-dBATCH",
            f"-sOutputFile={out_path}",
            str(in_path),
        ]
        subprocess.run(cmd, check=True, timeout=SUBPROC_TIMEOUT)
        return _stream_bytes(out_path.read_bytes(), "compressed.pdf", "application/pdf")


@app.post("/api/compress-bulk", tags=["PDF Operations"])
@limiter.limit(f"{RATE_LIMIT_PER_HOUR}/hour")
async def compress_pdfs_bulk(
    request: Request,
    files: List[UploadFile] = File(...),
    preset: str = Form("medium"),
):
    """Compress multiple PDF files and return as a zip archive."""
    async with SEMAPHORE:
        preset_map = {
            "low": "/screen",
            "medium": "/ebook",
            "high": "/prepress",
        }
        gs_preset = preset_map.get(preset, "/ebook")
        
        with tempfile.TemporaryDirectory() as td:
            out_zip = Path(td) / "compressed_pdfs.zip"
            
            with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                for i, file in enumerate(files):
                    try:
                        content = await file.read()
                        _ensure_size_limit(len(content))
                        
                        if not _is_pdf(content):
                            continue  # Skip non-PDF files
                        
                        # Create unique filename for this PDF
                        original_name = Path(file.filename or f"document_{i+1}.pdf").stem
                        compressed_name = f"compressed_{original_name}.pdf"
                        
                        # Compress the PDF
                        in_path = Path(td) / f"input_{i}.pdf"
                        out_path = Path(td) / f"output_{i}.pdf"
                        
                        in_path.write_bytes(content)
                        
                        cmd = [
                            "gs",
                            "-sDEVICE=pdfwrite",
                            "-dCompatibilityLevel=1.4",
                            f"-dPDFSETTINGS={gs_preset}",
                            "-dNOPAUSE",
                            "-dQUIET",
                            "-dBATCH",
                            f"-sOutputFile={out_path}",
                            str(in_path),
                        ]
                        
                        subprocess.run(cmd, check=True, timeout=SUBPROC_TIMEOUT)
                        
                        # Add compressed PDF to zip
                        zf.write(out_path, arcname=compressed_name)
                        
                    except Exception as e:
                        logger.error(f"Error compressing {file.filename}: {str(e)}")
                        continue
            
            # Return the zip file
            zip_data = out_zip.read_bytes()
            return _stream_bytes(zip_data, "compressed_pdfs.zip", "application/zip")


@app.post("/api/convert/pdf-to-images")
async def pdf_to_images(
    file: UploadFile = File(...),
    image_format: str = Form("png"),
    dpi: int = Form(150),
):
    pdf_bytes = await file.read()
    _ensure_size_limit(len(pdf_bytes))
    if not _is_pdf(pdf_bytes):
        return JSONResponse({"error": "Only PDF files are allowed"}, status_code=400)
    fmt = image_format.lower()
    if fmt not in ("png", "jpeg", "jpg"):
        fmt = "png"
    pil_format = "PNG" if fmt == "png" else "JPEG"
    images = convert_from_bytes(pdf_bytes, dpi=dpi)
    with tempfile.TemporaryDirectory() as td:
        base = Path(file.filename or "document.pdf").stem
        out_zip = Path(td) / f"{base}_images.zip"
        with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zf:
            for idx, img in enumerate(images, start=1):
                fn = f"{base}_page_{idx:03d}.{fmt if fmt != 'jpg' else 'jpg'}"
                p = Path(td) / fn
                params = {}
                if pil_format == "JPEG":
                    params["quality"] = 85
                    params["optimize"] = True
                img.save(p, pil_format, **params)
                zf.write(p, arcname=fn)
        return _stream_bytes(out_zip.read_bytes(), out_zip.name, "application/zip")


@app.post("/api/convert/images-to-pdf")
async def images_to_pdf(files: List[UploadFile] = File(...)):
    # Single file for single output, multiple files for zip
    if len(files) == 1:
        # Single image - convert directly to PDF
        content = await files[0].read()
        _ensure_size_limit(len(content))
        if not _is_image(content):
            return JSONResponse({"error": "Only image files are allowed"}, status_code=400)
        pdf_bytes = img2pdf.convert([content])
        original_name = Path(files[0].filename or "image.jpg").stem
        return _stream_bytes(pdf_bytes, f"{original_name}.pdf", "application/pdf")
    else:
        # Multiple images - output as zip
        with tempfile.TemporaryDirectory() as td:
            out_zip = Path(td) / "converted_images.zip"
            
            with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                for i, f in enumerate(files):
                    try:
                        content = await f.read()
                        _ensure_size_limit(len(content))
                        if not _is_image(content):
                            continue  # Skip non-image files
                        
                        # Convert each image to separate PDF
                        pdf_bytes = img2pdf.convert([content])
                        
                        # Save as temporary PDF file
                        temp_pdf = Path(td) / f"converted_{i+1:03d}.pdf"
                        temp_pdf.write_bytes(pdf_bytes)
                        
                        # Add to zip
                        original_name = Path(f.filename or f"image_{i+1}.jpg").stem
                        zf.write(temp_pdf, arcname=f"{original_name}.pdf")
                        
                    except Exception as e:
                        logger.error(f"Error converting {f.filename}: {str(e)}")
                        continue
            
            # Return zip file
            zip_data = out_zip.read_bytes()
            return _stream_bytes(zip_data, "converted_images.zip", "application/zip")


@app.post("/api/convert/auto")
async def convert_auto(
    target: str = Form(...),
    dpi: int = Form(150),
    file: Optional[UploadFile] = File(None),
    files: Optional[List[UploadFile]] = File(None),
):
    """Single endpoint that detects input type and converts to desired target.

    Supported routes:
    - PDF -> DOCX  (target=docx)
    - PDF -> images zip (target=png|jpeg, optional dpi)
    - DOCX -> PDF  (target=pdf)
    - images -> PDF (target=pdf) for one or more images
    """
    try:
        src_files: List[UploadFile] = list(files or ([] if file is None else [file]))
        if not src_files:
            return Response(status_code=400, content="No file(s) uploaded")

        # Normalized target and supported groups
        target_lc = (target or "").lower()
        image_targets = {"png", "jpeg", "jpg", "webp", "tiff", "bmp"}

        # If multiple images and target pdf -> images to pdf; or re-encode images as images (zip)
        if len(src_files) > 1:
            if target_lc == "pdf":
                image_bytes_list: List[bytes] = []
                for f in src_files:
                    content_type = (f.content_type or "").lower()
                    name = (f.filename or "").lower()
                    content = await f.read()
                    _ensure_size_limit(len(content))
                    if not (content_type.startswith("image/") or name.endswith(('.png','.jpg','.jpeg','.bmp','.tif','.tiff'))) or not _is_image(content):
                        return Response(status_code=400, content="All files must be images for images->pdf")
                    image_bytes_list.append(content)
                pdf_bytes = img2pdf.convert(image_bytes_list)
                return _stream_bytes(pdf_bytes, "images.pdf", "application/pdf")
            elif target_lc in image_targets:
                # Bulk image re-encode to the selected format
                fmt = "jpeg" if target_lc in {"jpeg", "jpg"} else target_lc
                pil_format = {
                    "png": "PNG",
                    "jpeg": "JPEG",
                    "webp": "WEBP",
                    "tiff": "TIFF",
                    "bmp": "BMP",
                }[fmt]
                with tempfile.TemporaryDirectory() as td:
                    out_zip = Path(td) / f"images_{fmt}.zip"
                    with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                        for idx, f in enumerate(src_files, start=1):
                            content = await f.read()
                            _ensure_size_limit(len(content))
                            img = Image.open(io.BytesIO(content))
                            out_name = f"image_{idx:03d}.{ 'jpg' if fmt == 'jpeg' else fmt }"
                            out_path = Path(td) / out_name
                            save_params = {}
                            if pil_format == "JPEG":
                                save_params["quality"] = 85
                                save_params["optimize"] = True
                            if pil_format == "WEBP":
                                save_params["quality"] = 85
                                save_params["method"] = 4
                            if pil_format == "TIFF":
                                save_params["compression"] = "tiff_lzw"
                            img.save(out_path, pil_format, **save_params)
                            zf.write(out_path, arcname=out_name)
                    return _stream_bytes(out_zip.read_bytes(), out_zip.name, "application/zip")
            # Multiple non-image files are not supported
            return Response(status_code=400, content="Multiple file upload is only supported for images->pdf (target=pdf)")

        # Single file flow
        single = src_files[0]
        filename = (single.filename or "").lower()
        ext = Path(filename).suffix
        content_type = (single.content_type or "").lower()
        
        logger.info(f"Processing file: {filename} ({len(str(single.size or 0))} bytes), target: {target_lc}")
        
        raw = await single.read()
        _ensure_size_limit(len(raw))
        logger.info(f"File read successfully, size: {len(raw)} bytes")

        # PDF sources
        if ext == ".pdf" or content_type == "application/pdf" or _is_pdf(raw):
            logger.info("Detected PDF source")
            if target_lc in image_targets:
                # PDF -> images zip
                pdf_bytes = raw
                fmt = "jpeg" if target_lc in {"jpeg", "jpg"} else target_lc
                pil_format = {
                    "png": "PNG",
                    "jpeg": "JPEG",
                    "webp": "WEBP",
                    "tiff": "TIFF",
                    "bmp": "BMP",
                }[fmt]
                # Memory optimization for large PDFs
                pdf_size_mb = len(pdf_bytes) / (1024 * 1024)
                if pdf_size_mb > 8:  # Large PDF optimization for 20MB limit
                    logger.info(f"Large PDF detected ({pdf_size_mb:.1f} MB), using memory-optimized conversion")
                    # Reduce DPI for large files to prevent memory issues
                    optimized_dpi = min(dpi, 80) if pdf_size_mb > 15 else min(dpi, 100)
                    logger.info(f"Optimized DPI: {optimized_dpi} (original: {dpi})")
                    
                    # Use poppler-utils directly for better memory management
                    with tempfile.TemporaryDirectory() as td:
                        pdf_path = Path(td) / "input.pdf"
                        pdf_path.write_bytes(pdf_bytes)
                        
                        # Get page count first
                        page_count_cmd = ["pdfinfo", str(pdf_path)]
                        page_info = subprocess.run(page_count_cmd, capture_output=True, text=True, timeout=30)
                        page_count = 1
                        for line in page_info.stdout.splitlines():
                            if line.startswith("Pages:"):
                                page_count = int(line.split(":")[1].strip())
                                break
                        
                        # Limit pages for very large PDFs to prevent excessive processing
                        max_pages = 100
                        if page_count > max_pages:
                            logger.warning(f"PDF has {page_count} pages, limiting to first {max_pages} pages for performance")
                            page_count = max_pages
                        
                        logger.info(f"Processing {page_count} pages with poppler-utils")
                        
                        # Process page by page to avoid memory issues
                        base = Path(filename or "document.pdf").stem
                        out_zip = Path(td) / f"{base}_images.zip"
                        
                        with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                            for page_num in range(1, page_count + 1):
                                try:
                                    # Convert single page using pdftoppm
                                    output_pattern = str(Path(td) / f"page_{page_num:03d}")
                                    
                                    # Fix pdftoppm command format
                                    if fmt == "jpeg":
                                        cmd = [
                                            "pdftoppm", 
                                            "-f", str(page_num), 
                                            "-l", str(page_num),
                                            "-r", str(optimized_dpi),
                                            "-jpeg",
                                            str(pdf_path), 
                                            output_pattern
                                        ]
                                    elif fmt == "png":
                                        cmd = [
                                            "pdftoppm", 
                                            "-f", str(page_num), 
                                            "-l", str(page_num),
                                            "-r", str(optimized_dpi),
                                            "-png",
                                            str(pdf_path), 
                                            output_pattern
                                        ]
                                    else:
                                        # Default to PNG for other formats
                                        cmd = [
                                            "pdftoppm", 
                                            "-f", str(page_num), 
                                            "-l", str(page_num),
                                            "-r", str(optimized_dpi),
                                            "-png",
                                            str(pdf_path), 
                                            output_pattern
                                        ]
                                    
                                    logger.info(f"Processing page {page_num}/{page_count}")
                                    subprocess.run(cmd, check=True, timeout=60, cwd=td)
                                    
                                    # Find the generated file
                                    generated_files = list(Path(td).glob(f"page_{page_num:03d}*"))
                                    if generated_files:
                                        img_file = generated_files[0]
                                        ext_out = 'jpg' if fmt == 'jpeg' else fmt
                                        fn = f"{base}_page_{page_num:03d}.{ext_out}"
                                        zf.write(img_file, arcname=fn)
                                        logger.info(f"Page {page_num} processed successfully")
                                    
                                except Exception as e:
                                    logger.error(f"Error processing page {page_num}: {e}")
                                    continue
                        
                        logger.info(f"Memory-optimized conversion completed: {out_zip}")
                        return _stream_bytes(out_zip.read_bytes(), out_zip.name, "application/zip")
                else:
                    # Standard conversion for smaller PDFs
                    logger.info(f"Converting PDF to {fmt} images with DPI {dpi}")
                    images = convert_from_bytes(pdf_bytes, dpi=dpi)
                    logger.info(f"PDF rendered to {len(images)} images")
                with tempfile.TemporaryDirectory() as td:
                    base = Path(filename or "document.pdf").stem
                    out_zip = Path(td) / f"{base}_images.zip"
                    with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                        for idx, img in enumerate(images, start=1):
                            ext_out = 'jpg' if fmt == 'jpeg' else fmt
                            fn = f"{base}_page_{idx:03d}.{ext_out}"
                            p = Path(td) / fn
                            params = {}
                            if pil_format == "JPEG":
                                params["quality"] = 85
                                params["optimize"] = True
                            if pil_format == "WEBP":
                                params["quality"] = 85
                                params["method"] = 4
                            if pil_format == "TIFF":
                                params["compression"] = "tiff_lzw"
                            img.save(p, pil_format, **params)
                            zf.write(p, arcname=fn)
                    logger.info(f"Images saved to zip: {out_zip}")
                    return _stream_bytes(out_zip.read_bytes(), out_zip.name, "application/zip")
            elif target_lc == "docx":
                # PDF -> DOCX (via pdftotext)
                logger.info("Converting PDF to DOCX")
                with tempfile.TemporaryDirectory() as td:
                    in_path = Path(td) / "input.pdf"
                    txt_path = Path(td) / "out.txt"
                    in_path.write_bytes(raw)
                    cmd = ["pdftotext", "-layout", str(in_path), str(txt_path)]
                    logger.info(f"Running pdftotext: {' '.join(cmd)}")
                    subprocess.run(cmd, check=True, timeout=SUBPROC_TIMEOUT)
                    text = txt_path.read_text(encoding="utf-8", errors="ignore")
                    doc = Document()
                    for line in text.splitlines():
                        doc.add_paragraph(line)
                    docx_path = Path(td) / "converted.docx"
                    doc.save(str(docx_path))
                    logger.info("DOCX conversion completed")
                    return _stream_bytes(
                        docx_path.read_bytes(),
                        "converted.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
            elif target_lc == "txt":
                # PDF -> TXT
                logger.info("Converting PDF to TXT")
                with tempfile.TemporaryDirectory() as td:
                    in_path = Path(td) / "input.pdf"
                    txt_path = Path(td) / "out.txt"
                    in_path.write_bytes(raw)
                    cmd = ["pdftotext", "-layout", str(in_path), str(txt_path)]
                    logger.info(f"Running pdftotext: {' '.join(cmd)}")
                    subprocess.run(cmd, check=True, timeout=SUBPROC_TIMEOUT)
                    logger.info("TXT conversion completed")
                    return _stream_bytes(
                        txt_path.read_bytes(),
                        "converted.txt",
                        "text/plain; charset=utf-8",
                    )
            else:
                return Response(status_code=400, content="Unsupported target for PDF. Use docx/png/jpeg.")

        # DOCX source
        if ext == ".docx" or content_type.endswith("officedocument.wordprocessingml.document") or _is_docx(raw, filename):
            logger.info("Detected DOCX source")
            if target_lc == "pdf":
                with tempfile.TemporaryDirectory() as td:
                    in_path = Path(td) / (single.filename or "input.docx")
                    pdf_out = Path(td)
                    in_path.write_bytes(raw)
                    cmd = [
                        "soffice",
                        "--headless",
                        "--norestore",
                        "--nolockcheck",
                        "--convert-to",
                        "pdf:writer_pdf_Export",
                        "--outdir",
                        str(pdf_out),
                        str(in_path),
                    ]
                    logger.info(f"Running LibreOffice conversion: {' '.join(cmd)}")
                    subprocess.run(cmd, check=True, timeout=SUBPROC_TIMEOUT)
                    out_pdf = pdf_out / (Path(in_path).stem + ".pdf")
                    logger.info("DOCX to PDF conversion completed")
                    return _stream_bytes(out_pdf.read_bytes(), out_pdf.name, "application/pdf")
            elif target_lc in image_targets:
                # DOCX -> images: first convert to PDF then render pages
                logger.info(f"Converting DOCX to {target_lc} images")
                with tempfile.TemporaryDirectory() as td:
                    in_path = Path(td) / (single.filename or "input.docx")
                    pdf_out = Path(td)
                    in_path.write_bytes(raw)
                    cmd = [
                        "soffice",
                        "--headless",
                        "--norestore",
                        "--nolockcheck",
                        "--convert-to",
                        "pdf:writer_pdf_Export",
                        "--outdir",
                        str(pdf_out),
                        str(in_path),
                    ]
                    logger.info(f"Running LibreOffice conversion: {' '.join(cmd)}")
                    subprocess.run(cmd, check=True, timeout=SUBPROC_TIMEOUT)
                    out_pdf = pdf_out / (Path(in_path).stem + ".pdf")
                    pdf_bytes = out_pdf.read_bytes()
                    fmt = "jpeg" if target_lc in {"jpeg", "jpg"} else target_lc
                    pil_format = {
                        "png": "PNG",
                        "jpeg": "JPEG",
                        "webp": "WEBP",
                        "tiff": "TIFF",
                        "bmp": "BMP",
                    }[fmt]
                    images = convert_from_bytes(pdf_bytes, dpi=dpi)
                    logger.info(f"DOCX rendered to {len(images)} images")
                    out_zip = Path(td) / f"docx_images.zip"
                    with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                        for idx, img in enumerate(images, start=1):
                            ext_out = 'jpg' if fmt == 'jpeg' else fmt
                            fn = f"page_{idx:03d}.{ext_out}"
                            p = Path(td) / fn
                            params = {}
                            if pil_format == "JPEG":
                                params["quality"] = 85
                                params["optimize"] = True
                            if pil_format == "WEBP":
                                params["quality"] = 85
                                params["method"] = 4
                            if pil_format == "TIFF":
                                params["compression"] = "tiff_lzw"
                            img.save(p, pil_format, **params)
                            zf.write(p, arcname=fn)
                    logger.info(f"DOCX images saved to zip: {out_zip}")
                    return _stream_bytes(out_zip.read_bytes(), out_zip.name, "application/zip")
            elif target_lc == "txt":
                # DOCX -> TXT via python-docx
                logger.info("Converting DOCX to TXT")
                doc = Document(io.BytesIO(raw))
                text = "\n".join(p.text for p in doc.paragraphs)
                logger.info("DOCX to TXT conversion completed")
                return _stream_bytes(text.encode("utf-8"), "converted.txt", "text/plain; charset=utf-8")
            else:
                return Response(status_code=400, content="Unsupported target for DOCX. Use pdf.")

        # Image source(s)
        if content_type.startswith("image/") or ext in (".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff") or _is_image(raw):
            logger.info("Detected image source")
            if target_lc == "pdf":
                image_bytes_list = [raw]
                pdf_bytes = img2pdf.convert(image_bytes_list)
                logger.info("Image to PDF conversion completed")
                return _stream_bytes(pdf_bytes, "image.pdf", "application/pdf")
            elif target_lc in image_targets:
                # Re-encode single image
                fmt = "jpeg" if target_lc in {"jpeg", "jpg"} else target_lc
                pil_format = {
                    "png": "PNG",
                    "jpeg": "JPEG",
                    "webp": "WEBP",
                    "tiff": "TIFF",
                    "bmp": "BMP",
                }[fmt]
                img = Image.open(io.BytesIO(raw))
                buf = io.BytesIO()
                save_params = {}
                if pil_format == "JPEG":
                    save_params["quality"] = 85
                    save_params["optimize"] = True
                if pil_format == "WEBP":
                    save_params["quality"] = 85
                    save_params["method"] = 4
                if pil_format == "TIFF":
                    save_params["compression"] = "tiff_lzw"
                img.save(buf, pil_format, **save_params)
                ext_out = 'jpg' if fmt == 'jpeg' else fmt
                media = f"image/{'jpeg' if fmt == 'jpeg' else ext_out}"
                logger.info(f"Image re-encoded to {fmt}")
                return _stream_bytes(buf.getvalue(), f"image.{ext_out}", media)
            else:
                return Response(status_code=400, content="Unsupported target for images. Use pdf.")

        return Response(status_code=400, content="Unsupported source/target combination")
        
    except subprocess.TimeoutExpired as e:
        logger.error(f"Subprocess timeout: {e}")
        return JSONResponse({"error": f"Conversion timed out after {SUBPROC_TIMEOUT} seconds. File may be too large or complex."}, status_code=408)
    except subprocess.CalledProcessError as e:
        logger.error(f"Subprocess error: {e}")
        return JSONResponse({"error": f"Conversion failed: {e}"}, status_code=500)
    except ValueError as e:
        logger.error(f"Validation error: {e}")
        return JSONResponse({"error": str(e)}, status_code=400)
    except Exception as e:
        logger.error(f"Unexpected error in convert_auto: {e}", exc_info=True)
        return JSONResponse({"error": "Internal server error during conversion"}, status_code=500)

@app.post("/api/convert/pdf-to-docx")
async def pdf_to_docx(file: UploadFile = File(...)):
    # Extract text with poppler's pdftotext, then place into docx
    with tempfile.TemporaryDirectory() as td:
        in_path = Path(td) / "input.pdf"
        txt_path = Path(td) / "out.txt"
        in_path.write_bytes(await file.read())
        cmd = ["pdftotext", "-layout", str(in_path), str(txt_path)]
        subprocess.run(cmd, check=True)

        text = txt_path.read_text(encoding="utf-8", errors="ignore")
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        docx_path = Path(td) / "converted.docx"
        doc.save(str(docx_path))
        return _stream_bytes(docx_path.read_bytes(), "converted.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.post("/api/convert/docx-to-pdf")
async def docx_to_pdf(file: UploadFile = File(...)):
    with tempfile.TemporaryDirectory() as td:
        in_path = Path(td) / (file.filename or "input.docx")
        pdf_out = Path(td)
        in_path.write_bytes(await file.read())
        # Use LibreOffice headless conversion
        cmd = [
            "soffice",
            "--headless",
            "--norestore",
            "--nolockcheck",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(pdf_out),
            str(in_path),
        ]
        subprocess.run(cmd, check=True)
        out_pdf = pdf_out / (Path(in_path).stem + ".pdf")
        return _stream_bytes(out_pdf.read_bytes(), out_pdf.name, "application/pdf")


@app.post("/api/ocr")
async def ocr_file(
    file: UploadFile = File(...),
    lang: str = Form("tur+eng"),
    output: str = Form("txt"),  # txt or pdf
    dpi: int = Form(200),
):
    name = (file.filename or "document").lower()
    ext = Path(name).suffix
    content = await file.read()
    with tempfile.TemporaryDirectory() as td:
        if ext in [".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"]:
            img_path = Path(td) / f"input{ext}"
            img_path.write_bytes(content)
            if output == "pdf":
                page_pdf = Path(td) / "ocr.pdf"
                cmd = ["tesseract", str(img_path), str(page_pdf.with_suffix("")), "-l", lang, "pdf"]
                subprocess.run(cmd, check=True)
                return _stream_bytes(page_pdf.read_bytes(), "ocr.pdf", "application/pdf")
            else:
                # text
                cmd = ["tesseract", str(img_path), "stdout", "-l", lang]
                res = subprocess.run(cmd, check=True, capture_output=True)
                txt = res.stdout.decode("utf-8", errors="ignore")
                return _stream_bytes(txt.encode("utf-8"), "ocr.txt", "text/plain; charset=utf-8")
        else:
            # Assume PDF: rasterize then OCR each page
            pdf_path = Path(td) / "input.pdf"
            pdf_path.write_bytes(content)
            # Render with pdftoppm to PNG files
            img_prefix = Path(td) / "page"
            cmd = ["pdftoppm", "-r", str(dpi), "-png", str(pdf_path), str(img_prefix)]
            subprocess.run(cmd, check=True)
            images = sorted(Path(td).glob("page-*.png"))
            if output == "pdf":
                # OCR each page to PDF then merge
                page_pdfs = []
                for idx, img in enumerate(images, start=1):
                    out_base = Path(td) / f"ocr_page_{idx:03d}"
                    cmd = ["tesseract", str(img), str(out_base), "-l", lang, "pdf"]
                    subprocess.run(cmd, check=True)
                    page_pdfs.append(out_base.with_suffix(".pdf"))
                writer = PdfWriter()
                for p in page_pdfs:
                    reader = PdfReader(str(p))
                    for page in reader.pages:
                        writer.add_page(page)
                buf = io.BytesIO()
                writer.write(buf)
                writer.close()
                buf.seek(0)
                return _stream_bytes(buf.getvalue(), "ocr.pdf", "application/pdf")
            else:
                # txt: concatenate pages
                combined = []
                for img in images:
                    cmd = ["tesseract", str(img), "stdout", "-l", lang]
                    res = subprocess.run(cmd, check=True, capture_output=True)
                    combined.append(res.stdout.decode("utf-8", errors="ignore"))
                text_all = "\n".join(combined)
                return _stream_bytes(text_all.encode("utf-8"), "ocr.txt", "text/plain; charset=utf-8")


@app.post("/api/ocr-bulk", tags=["PDF Operations"])
@limiter.limit(f"{RATE_LIMIT_PER_HOUR}/hour")
async def ocr_files_bulk(
    request: Request,
    files: List[UploadFile] = File(...),
    lang: str = Form("tur+eng"),
    output: str = Form("txt"),  # txt or pdf
    dpi: int = Form(200),
):
    """OCR multiple files and return as a zip archive."""
    async with SEMAPHORE:
        # Single file for single output, multiple files for zip
        if len(files) == 1:
            # Single file - use existing OCR endpoint
            return await ocr_file(files[0], lang, output, dpi)
        else:
            # Multiple files - output as zip
            with tempfile.TemporaryDirectory() as td:
                out_zip = Path(td) / "ocr_results.zip"
                
                with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                    for i, file in enumerate(files):
                        try:
                            name = (file.filename or f"document_{i+1}").lower()
                            ext = Path(name).suffix
                            content = await file.read()
                            
                            if ext in [".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"]:
                                # Image file
                                img_path = Path(td) / f"input_{i}{ext}"
                                img_path.write_bytes(content)
                                
                                if output == "pdf":
                                    page_pdf = Path(td) / f"ocr_{i+1:03d}.pdf"
                                    cmd = ["tesseract", str(img_path), str(page_pdf.with_suffix("")), "-l", lang, "pdf"]
                                    subprocess.run(cmd, check=True)
                                    zf.write(page_pdf, arcname=f"ocr_{Path(name).stem}.pdf")
                                else:
                                    # text
                                    cmd = ["tesseract", str(img_path), "stdout", "-l", lang]
                                    res = subprocess.run(cmd, check=True, capture_output=True)
                                    txt = res.stdout.decode("utf-8", errors="ignore")
                                    txt_path = Path(td) / f"ocr_{i+1:03d}.txt"
                                    txt_path.write_text(txt, encoding="utf-8")
                                    zf.write(txt_path, arcname=f"ocr_{Path(name).stem}.txt")
                                    
                            else:
                                # PDF file
                                pdf_path = Path(td) / f"input_{i}.pdf"
                                pdf_path.write_bytes(content)
                                
                                # Render with pdftoppm to PNG files
                                img_prefix = Path(td) / f"page_{i}"
                                cmd = ["pdftoppm", "-r", str(dpi), "-png", str(pdf_path), str(img_prefix)]
                                subprocess.run(cmd, check=True)
                                images = sorted(Path(td).glob(f"page_{i}-*.png"))
                                
                                if output == "pdf":
                                    # OCR each page to PDF then merge
                                    page_pdfs = []
                                    for idx, img in enumerate(images, start=1):
                                        out_base = Path(td) / f"ocr_{i+1:03d}_page_{idx:03d}"
                                        cmd = ["tesseract", str(img), str(out_base), "-l", lang, "pdf"]
                                        subprocess.run(cmd, check=True)
                                        page_pdfs.append(out_base.with_suffix(".pdf"))
                                    
                                    writer = PdfWriter()
                                    for p in page_pdfs:
                                        reader = PdfReader(str(p))
                                        for page in reader.pages:
                                            writer.add_page(page)
                                    
                                    buf = io.BytesIO()
                                    writer.write(buf)
                                    writer.close()
                                    buf.seek(0)
                                    
                                    temp_pdf = Path(td) / f"ocr_{i+1:03d}.pdf"
                                    temp_pdf.write_bytes(buf.getvalue())
                                    zf.write(temp_pdf, arcname=f"ocr_{Path(name).stem}.pdf")
                                    
                                else:
                                    # txt: concatenate pages
                                    combined = []
                                    for img in images:
                                        cmd = ["tesseract", str(img), "stdout", "-l", lang]
                                        res = subprocess.run(cmd, check=True, capture_output=True)
                                        combined.append(res.stdout.decode("utf-8", errors="ignore"))
                                    text_all = "\n".join(combined)
                                    
                                    txt_path = Path(td) / f"ocr_{i+1:03d}.txt"
                                    txt_path.write_text(text_all, encoding="utf-8")
                                    zf.write(txt_path, arcname=f"ocr_{Path(name).stem}.txt")
                                    
                        except Exception as e:
                            logger.error(f"Error processing {file.filename}: {str(e)}")
                            continue
                
                # Return zip file
                zip_data = out_zip.read_bytes()
                return _stream_bytes(zip_data, "ocr_results.zip", "application/zip")


@app.post("/api/encrypt")
async def encrypt_pdf(
    file: UploadFile = File(...),
    user_password: str = Form(...),
    owner_password: Optional[str] = Form(None),
    allow_print: bool = Form(True),
    allow_modify: bool = Form(False),
    allow_copy: bool = Form(False),
):
    try:
        logger.info(f"Encrypt request received - File: {file.filename}, Size: {file.size}")
        logger.info(f"Password: {user_password[:3]}***, Owner: {owner_password[:3] if owner_password else 'None'}***")
        logger.info(f"Permissions - Print: {allow_print}, Modify: {allow_modify}, Copy: {allow_copy}")
        
        with tempfile.TemporaryDirectory() as td:
            in_path = Path(td) / "input.pdf"
            out_path = Path(td) / "encrypted.pdf"
            
            # Read and save file
            content = await file.read()
            logger.info(f"File content read, size: {len(content)} bytes")
            in_path.write_bytes(content)
            
            # Open PDF
            logger.info("Opening PDF with pikepdf...")
            with pikepdf.open(str(in_path)) as pdf:
                logger.info(f"PDF opened successfully, pages: {len(pdf.pages)}")
                
                # Set permissions - latest API for pikepdf 9.0.0
                try:
                    # Latest API (pikepdf 9.0.0)
                    perms = pikepdf.Permissions(
                        print_document=allow_print,
                        modify_annotation=allow_modify,
                        extract=allow_copy,
                        annotate=True,
                    )
                    logger.info(f"Permissions set (pikepdf 9.0.0): {perms}")
                except Exception as perm_error:
                    logger.warning(f"Permissions error: {perm_error}, using default permissions")
                    # Default permissions
                    perms = pikepdf.Permissions()
                    logger.info("Using default permissions")
                
                # Save encrypted
                logger.info("Saving encrypted PDF...")
                pdf.save(
                    str(out_path),
                    encryption=pikepdf.Encryption(
                        user=user_password,
                        owner=owner_password or user_password,
                        R=6,  # AES-256
                        allow=perms,
                    ),
                )
                logger.info("PDF encrypted successfully")
            
            # Return encrypted file
            result_bytes = out_path.read_bytes()
            logger.info(f"Returning encrypted file, size: {len(result_bytes)} bytes")
            return _stream_bytes(result_bytes, "encrypted.pdf", "application/pdf")
            
    except Exception as e:
        logger.error(f"Encrypt error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Encryption failed: {str(e)}")


@app.post("/api/decrypt")
async def decrypt_pdf(file: UploadFile = File(...), password: str = Form(...)):
    with tempfile.TemporaryDirectory() as td:
        in_path = Path(td) / "input.pdf"
        out_path = Path(td) / "decrypted.pdf"
        in_path.write_bytes(await file.read())
        with pikepdf.open(str(in_path), password=password) as pdf:
            pdf.save(str(out_path))
        return _stream_bytes(out_path.read_bytes(), "decrypted.pdf", "application/pdf")


@app.post("/api/decrypt-bulk", tags=["PDF Operations"])
@limiter.limit(f"{RATE_LIMIT_PER_HOUR}/hour")
async def decrypt_pdfs_bulk(
    request: Request,
    files: List[UploadFile] = File(...),
    password: str = Form(...),
):
    """Decrypt multiple PDF files and return as a zip archive."""
    async with SEMAPHORE:
        # Single file for single output, multiple files for zip
        if len(files) == 1:
            # Single file - use existing decrypt endpoint
            return await decrypt_pdf(files[0], password)
        else:
            # Multiple files - output as zip
            with tempfile.TemporaryDirectory() as td:
                out_zip = Path(td) / "decrypted_pdfs.zip"
                
                with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                    for i, file in enumerate(files):
                        try:
                            content = await file.read()
                            in_path = Path(td) / f"input_{i}.pdf"
                            out_path = Path(td) / f"decrypted_{i}.pdf"
                            
                            in_path.write_bytes(content)
                            
                            with pikepdf.open(str(in_path), password=password) as pdf:
                                pdf.save(str(out_path))
                                
                                # Add to zip
                                original_name = Path(file.filename or f"document_{i+1}.pdf").stem
                                zf.write(out_path, arcname=f"decrypted_{original_name}.pdf")
                                
                        except Exception as e:
                            logger.error(f"Error decrypting {file.filename}: {str(e)}")
                            continue
                
                # Return zip file
                zip_data = out_zip.read_bytes()
                return _stream_bytes(zip_data, "decrypted_pdfs.zip", "application/zip")




@app.get("/health")
async def health_check():
    """Health check endpoint for monitoring and load balancers."""
    try:
        # Basic system checks
        import psutil
        memory = psutil.virtual_memory()
        disk = psutil.disk_usage('/')
        
        return {
            "status": "healthy",
            "timestamp": datetime.now().isoformat(),
            "uptime_seconds": int(time.time() - SERVICE_START_TIME),
            "environment": os.getenv("ENVIRONMENT", "development"),
            "version": "2.0.0",
            "system": {
                "memory_percent": memory.percent,
                "disk_percent": disk.percent,
                "cpu_count": psutil.cpu_count()
            }
        }
    except Exception as e:
        logger.error(f"Health check failed: {str(e)}")
        return JSONResponse(
            status_code=503,
            content={
                "status": "unhealthy",
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            }
        )

@app.get("/metrics")
async def metrics():
    """Basic metrics endpoint for monitoring."""
    try:
        import psutil
        memory = psutil.virtual_memory()
        disk = psutil.disk_usage('/')
        
        return {
            "memory_used_mb": memory.used // (1024 * 1024),
            "memory_total_mb": memory.total // (1024 * 1024),
            "memory_percent": memory.percent,
            "disk_used_gb": disk.used // (1024 * 1024 * 1024),
            "disk_total_gb": disk.total // (1024 * 1024 * 1024),
            "disk_percent": disk.percent,
            "uptime_seconds": int(time.time() - SERVICE_START_TIME)
        }
    except Exception as e:
        logger.error(f"Metrics failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Metrics collection failed")

@app.get("/favicon.ico")
def get_favicon():
	from PIL import Image
	buf = io.BytesIO()
	img = Image.new("RGBA", (64, 64), (59, 130, 246, 255))  # Tailwind primary blue
	img.save(buf, format="ICO", sizes=[(16,16),(32,32),(64,64)])
	buf.seek(0)
	return Response(content=buf.getvalue(), media_type="image/x-icon")


@app.get("/.well-known/appspecific/com.chrome.devtools.json")
def well_known_devtools():
	return JSONResponse({"status": "ok", "demo": True})


